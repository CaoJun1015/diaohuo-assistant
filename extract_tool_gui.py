"""
联想配置提取工具 - 图形界面版本
支持拖拽操作，自动提取联想电脑配置
"""

import sys
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from PyQt6.QtWidgets import (QApplication, QMainWindow, QLabel, QVBoxLayout, 
                             QHBoxLayout, QWidget, QMessageBox, QProgressBar,
                             QFrame)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize
from PyQt6.QtGui import QFont, QDragEnterEvent, QDropEvent, QIcon
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

LENOVO_PREFIXES = [
    "小新Pro16GT-ultra", "小新Pro16GT-", "小新Pro16C-", "小新Pro16-",
    "小新Pro14C-", "小新Pro14GT-", "小新Pro14-",
    "小新SE-16C-", "小新SE-14C-", "小新SE-16", "小新SE-14",
    "小新SE-16C", "小新SE-14C",
    "小新Air14-", "小新Air15-", "小新Air14", "小新Air15",
    "Y9000P-16", "Y9000P", "Y7000P", "Y7000",
    "R9000P", "R9000", "r9000p", "r9000",
    "拯救者", "LEGION",
    "ThinkPad", "ThinkBook", "Think",
    "来酷", "ideapad", "IdeaPad",
]

CPU_PATTERN = re.compile(
    r'(I[3579][-\s]?\d{3,5}[A-Za-z]*|R[579][-\s]?\d{3,4}[A-Za-z]*'
    r'|U\d+[-\s]?\d*[A-Za-z]*|ultra\d+[-\s]?\d*[A-Za-z]*)',
    re.IGNORECASE
)


class ExtractThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(bool, str, str)
    
    def __init__(self, input_file):
        super().__init__()
        self.input_file = input_file
        self.products = []
    
    def run(self):
        try:
            self.progress.emit(10)
            self.products = self.extract_products(self.input_file)
            self.progress.emit(50)
            
            if not self.products:
                self.finished.emit(False, "未找到联想电脑配置", "")
                return
            
            output_file = self.create_output(self.products)
            self.progress.emit(100)
            self.finished.emit(True, f"成功提取 {len(self.products)} 条配置", output_file)
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            self.finished.emit(False, f"处理失败: {str(e)}", "")
    
    def extract_products(self, filepath):
        raw_segments = []
        
        try:
            doc = Document(filepath)
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts = re.split(r'[ \t]{3,}', text)
                    for p in parts:
                        p = p.strip()
                        if p and len(p) > 3:
                            raw_segments.append(p)
            
            if not raw_segments and doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text:
                                row_text.append(text)
                        if row_text:
                            combined = '    '.join(row_text)
                            raw_segments.append(combined)
            
            if raw_segments:
                return self.parse_products(raw_segments)
        except:
            pass
        
        try:
            with zipfile.ZipFile(filepath, 'r') as zf:
                if 'word/document.xml' not in zf.namelist():
                    return []
                
                with zf.open('word/document.xml') as doc_xml:
                    content = doc_xml.read()
                    root = ET.fromstring(content)
                    
                    namespaces = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    }
                    
                    for para in root.findall('.//w:p', namespaces):
                        para_text = []
                        for text_elem in para.findall('.//w:t', namespaces):
                            if text_elem.text:
                                para_text.append(text_elem.text)
                        
                        text = ''.join(para_text).strip()
                        if text and len(text) > 3:
                            parts = re.split(r'[ \t]{3,}', text)
                            for p in parts:
                                p = p.strip()
                                if p:
                                    raw_segments.append(p)
                    
                    return self.parse_products(raw_segments)
        except:
            return []
    
    def parse_products(self, raw_segments):
        lenovo_segments = [seg for seg in raw_segments if any(prefix.lower() in seg.lower() for prefix in LENOVO_PREFIXES)]
        
        products = []
        for seg in lenovo_segments:
            product = self.parse_single(seg)
            if product and product['series'] and product['cpu']:
                products.append(product)
        
        unique_products = []
        seen = set()
        for p in products:
            key = (p['series'], p['cpu'], p['ram'], p['storage'])
            if key not in seen:
                seen.add(key)
                unique_products.append(p)
        
        return unique_products
    
    def parse_single(self, text):
        result = {'series': '', 'cpu': '', 'ram': '', 'storage': '', 'gpu': '', 'screen': '', 'note': ''}
        
        raw = text.strip()
        if not raw:
            return None
        
        for prefix in LENOVO_PREFIXES:
            if prefix.lower() in raw.lower():
                result['series'] = prefix
                break
        
        if not result['series']:
            m = re.match(r'^([\u4e00-\u9fa5A-Za-z][\u4e00-\u9fa5A-Za-z0-9\-]*)', raw)
            if m and len(m.group(1)) >= 2:
                result['series'] = m.group(1)
        
        if not result['series']:
            return None
        
        result['series'] = result['series'].replace("联想", "").strip()
        
        cpu_m = CPU_PATTERN.search(raw)
        if cpu_m:
            result['cpu'] = cpu_m.group(0).strip().upper()
        
        remainder = re.sub(r'\b8G\s+(5060|5070|4060|4070)', 'VRAM', raw)
        ram_m = re.search(r'\b(\d+)\s*[Gg][Bb]?\b', remainder)
        if ram_m:
            result['ram'] = ram_m.group(1) + "G"
        
        storage_m = re.search(r'\b(\d+)\s*[Tt][Bb]?\b', remainder)
        if storage_m:
            result['storage'] = storage_m.group(1) + "T"
        else:
            storage_m = re.search(r'\b(\d+)\s*[Gg][Bb]?\b', remainder)
            if storage_m:
                result['storage'] = storage_m.group(1) + "G"
        
        gpu_m = re.search(r'(\d+)\s*([45]0\d{1,2})', remainder)
        if gpu_m:
            result['gpu'] = gpu_m.group(1) + gpu_m.group(2)
        elif re.search(r'(RTX\s*\d+|GTX\s*\d+)', remainder, re.IGNORECASE):
            gpu_m = re.search(r'(RTX\s*\d+|GTX\s*\d+)', remainder, re.IGNORECASE)
            result['gpu'] = gpu_m.group(0).strip()
        elif re.search(r'(集成|集显)', remainder):
            result['gpu'] = "集成"
        
        screen_m = re.search(r'\b(\d{2}(?:\.\d)?)\s*(?:英寸|寸|屏)?', remainder)
        if screen_m:
            val = screen_m.group(1)
            if 10 <= float(val) <= 18:
                result['screen'] = val
        
        note_parts = []
        colors = ['碳晶灰', '鸽子灰', '银灰', '黑色', '白色', '蓝色', '银色', '深空灰']
        for color in colors:
            if color in raw:
                note_parts.append(color)
                break
        
        if '新品' in raw:
            note_parts.append('新品')
        if 'Win11' in raw or 'win11' in raw:
            note_parts.append('Win11')
        
        if note_parts:
            result['note'] = ' '.join(note_parts)
        
        return result
    
    def create_output(self, products):
        app_path = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(__file__)
        input_name = os.path.splitext(os.path.basename(self.input_file))[0]
        output_file = os.path.join(app_path, f"{input_name}_联想配置.docx")
        
        doc = Document()
        
        title = doc.add_heading('联想电脑配置清单', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info_run = info_para.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        info_run.font.size = Pt(9)
        info_run.font.italic = True
        
        doc.add_paragraph()
        
        intro = doc.add_paragraph()
        intro_run = intro.add_run('说明: 本文档按照调货助手数据库格式生成，可直接导入使用。')
        intro_run.font.size = Pt(10)
        intro_run.font.italic = True
        
        doc.add_paragraph()
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_cells = table.rows[0].cells
        headers = ['系列', 'CPU', '内存', '硬盘', '显卡', '屏幕', '备注']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, width in enumerate([Cm(4), Cm(3.5), Cm(2), Cm(2), Cm(2), Cm(1.5), Cm(3)]):
            for cell in table.columns[i].cells:
                cell.width = width
        
        for product in products:
            row_cells = table.add_row().cells
            row_cells[0].text = product.get('series', '')
            row_cells[1].text = product.get('cpu', '')
            row_cells[2].text = product.get('ram', '')
            row_cells[3].text = product.get('storage', '')
            row_cells[4].text = product.get('gpu', '')
            row_cells[5].text = product.get('screen', '')
            row_cells[6].text = product.get('note', '')
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        doc.add_paragraph()
        
        stats = doc.add_paragraph()
        stats.alignment = WD_ALIGN_PARAGRAPH.LEFT
        stats_run = stats.add_run(f'共提取 {len(products)} 条联想电脑配置')
        stats_run.font.size = Pt(11)
        stats_run.font.bold = True
        
        doc.save(output_file)
        return output_file


class DropArea(QFrame):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setFrameStyle(QFrame.Shape.Box | QFrame.Shadow.Raised)
        self.setLineWidth(2)
        self.setMinimumSize(400, 300)
        
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.icon_label = QLabel("📄")
        self.icon_label.setFont(QFont("Arial", 60))
        self.icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.text_label = QLabel("拖拽Word文档到此处")
        self.text_label.setFont(QFont("Microsoft YaHei", 14))
        self.text_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.hint_label = QLabel("(支持WPS和微软Office)")
        self.hint_label.setFont(QFont("Microsoft YaHei", 10))
        self.hint_label.setStyleSheet("color: gray;")
        self.hint_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.format_label = QLabel("支持格式: .docx")
        self.format_label.setFont(QFont("Microsoft YaHei", 9))
        self.format_label.setStyleSheet("color: #888;")
        self.format_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        layout.addWidget(self.icon_label)
        layout.addWidget(self.text_label)
        layout.addWidget(self.hint_label)
        layout.addWidget(self.format_label)
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()
            if urls and urls[0].toLocalFile().lower().endswith('.docx'):
                event.acceptProposedAction()
                self.setStyleSheet("border: 3px dashed #4CAF50; background-color: #E8F5E9;")
                return
        event.ignore()
    
    def dragLeaveEvent(self, event):
        self.setStyleSheet("")
    
    def dropEvent(self, event: QDropEvent):
        self.setStyleSheet("")
        urls = event.mimeData().urls()
        if urls:
            file_path = urls[0].toLocalFile()
            if file_path.lower().endswith('.docx'):
                self.parent().process_file(file_path)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("联想配置提取工具")
        self.setMinimumSize(500, 400)
        self.thread = None
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        title_label = QLabel("联想配置提取工具")
        title_label.setFont(QFont("Microsoft YaHei", 18, QFont.Weight.Bold))
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(title_label)
        
        main_layout.addSpacing(20)
        
        self.drop_area = DropArea(self)
        main_layout.addWidget(self.drop_area, 1)
        
        main_layout.addSpacing(20)
        
        status_frame = QFrame()
        status_frame.setFrameStyle(QFrame.Shape.Box | QFrame.Shadow.Sunken)
        status_layout = QVBoxLayout(status_frame)
        
        self.status_label = QLabel("状态: 等待中...")
        self.status_label.setFont(QFont("Microsoft YaHei", 10))
        status_layout.addWidget(self.status_label)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        status_layout.addWidget(self.progress_bar)
        
        self.file_label = QLabel("")
        self.file_label.setFont(QFont("Microsoft YaHei", 9))
        self.file_label.setStyleSheet("color: #666;")
        self.file_label.setWordWrap(True)
        self.file_label.setVisible(False)
        status_layout.addWidget(self.file_label)
        
        main_layout.addWidget(status_frame)
    
    def process_file(self, file_path):
        self.file_label.setText(f"已接收: {os.path.basename(file_path)}")
        self.file_label.setVisible(True)
        self.status_label.setText("状态: 正在处理...")
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        self.thread = ExtractThread(file_path)
        self.thread.progress.connect(self.on_progress)
        self.thread.finished.connect(self.on_finished)
        self.thread.start()
    
    def on_progress(self, value):
        self.progress_bar.setValue(value)
    
    def on_finished(self, success, message, output_file):
        self.progress_bar.setVisible(False)
        
        if success:
            self.status_label.setText(f"✓ {message}")
            self.status_label.setStyleSheet("color: #4CAF50; font-weight: bold;")
            
            msg = QMessageBox(self)
            msg.setWindowTitle("处理完成")
            msg.setText(f"成功提取联想电脑配置！")
            msg.setInformativeText(f"输出文件:\n{output_file}")
            msg.setIcon(QMessageBox.Icon.Information)
            
            open_btn = msg.addButton("打开文件", QMessageBox.ButtonRole.AcceptRole)
            msg.addButton("确定", QMessageBox.ButtonRole.RejectRole)
            
            msg.exec()
            
            if msg.clickedButton() == open_btn:
                os.startfile(output_file)
            
            self.status_label.setText("状态: 等待中...")
            self.status_label.setStyleSheet("")
            self.file_label.setVisible(False)
        else:
            self.status_label.setText(f"✗ {message}")
            self.status_label.setStyleSheet("color: #F44336;")
            
            QMessageBox.warning(self, "处理失败", message)


def main():
    app = QApplication(sys.argv)
    app.setFont(QFont("Microsoft YaHei", 10))
    
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
