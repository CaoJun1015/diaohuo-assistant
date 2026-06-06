"""
Word 解析模块：读取上游价格表 docx，提取机型列表
重写版本：更稳健的按前缀分割与逐字段提取
支持WPS和微软Office格式
"""

import re
import zipfile
import xml.etree.ElementTree as ET
from docx import Document

# 已知机型前缀列表（按匹配优先级排列，越长越优先）
KNOWN_PREFIXES = [
    "小新Pro16GT-ultra", "小新Pro16GT-", "小新Pro16C-", "小新Pro16-",
    "小新Pro14C-", "小新Pro14GT-", "小新Pro14-",
    "小新SE-16C-", "小新SE-14C-", "小新SE-16", "小新SE-14",
    "小新SE-16C", "小新SE-14C",
    "Y9000P-16", "Y9000P", "Y7000P", "Y7000",
    "R9000P", "R9000", "r9000p", "r9000",
    "联想来酷", "来酷",
    "拯救者",
    "雷神E1", "雷神14英寸", "雷神",
    "ThinkPad", "ThinkBook", "Think",
]

# CPU 正则（大小写通吃）
CPU_PATTERN = re.compile(
    r'(I[3579][-\s]?\d{3,5}[A-Za-z]*|R[579][-\s]?\d{3,4}[A-Za-z]*'
    r'|U\d+[-\s]?\d*[A-Za-z]*|ultra\d+[-\s]?\d*[A-Za-z]*)',
    re.IGNORECASE
)


def parse_word_pricelist(filepath):
    """
    解析上游价格 Word 文档，返回机型字典列表。
    支持WPS和微软Office格式。
    """
    # 先尝试标准python-docx解析
    try:
        doc = Document(filepath)
        raw_segments = extract_text_from_docx(doc)
    except Exception as e:
        # 如果标准方法失败，尝试直接解析XML
        print(f"标准解析失败，尝试兼容模式: {e}")
        try:
            raw_segments = extract_text_from_xml(filepath)
        except Exception as e2:
            print(f"兼容模式也失败: {e2}")
            return []
    
    # 合并碎片：如果一个片段极短（<8字符）且看起来是上一段的尾部，则合并
    merged = []
    for seg in raw_segments:
        if merged and _is_tail_fragment(seg):
            merged[-1] += " " + seg
        else:
            merged.append(seg)

    # 对每个 merged 段，再按机型前缀分割（一段可能包含多条）
    records = []
    for seg in merged:
        sub = _split_by_prefix(seg)
        records.extend(sub)

    # 逐条解析
    products = []
    for rec in records:
        if len(rec) < 6:  # 太短不可能是一条完整记录
            continue
        parsed = _parse_single_record(rec)
        if parsed and parsed.get("series") and parsed["series"] != "未识别":
            products.append(parsed)

    return products


def extract_text_from_docx(doc):
    """从python-docx Document对象提取文本"""
    raw_segments = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        # 按两个以上连续空格或制表符拆分
        parts = re.split(r'[ \t]{3,}', text)
        for p in parts:
            p = p.strip()
            if p:
                raw_segments.append(p)
    
    # 如果没有段落，尝试从表格提取
    if not raw_segments and doc.tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts = re.split(r'[ \t]{3,}', text)
                        for p in parts:
                            p = p.strip()
                            if p:
                                raw_segments.append(p)
    
    return raw_segments


def extract_text_from_xml(filepath):
    """直接从docx的XML结构提取文本（兼容WPS）"""
    raw_segments = []
    
    # docx文件实际上是ZIP压缩包
    with zipfile.ZipFile(filepath, 'r') as zf:
        # 读取document.xml
        if 'word/document.xml' not in zf.namelist():
            raise Exception("无法找到document.xml")
        
        with zf.open('word/document.xml') as doc_xml:
            content = doc_xml.read()
            
            # 解析XML
            root = ET.fromstring(content)
            
            # 定义命名空间
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'wps': 'http://schemas.wps.cn/officeDocument/2006/main'  # WPS命名空间
            }
            
            # 查找所有文本节点
            for para in root.findall('.//w:p', namespaces):
                para_text = []
                for text_elem in para.findall('.//w:t', namespaces):
                    if text_elem.text:
                        para_text.append(text_elem.text)
                
                text = ''.join(para_text).strip()
                if text:
                    parts = re.split(r'[ \t]{3,}', text)
                    for p in parts:
                        p = p.strip()
                        if p:
                            raw_segments.append(p)
            
            # 如果没有找到段落，尝试表格
            if not raw_segments:
                for table in root.findall('.//w:tbl', namespaces):
                    for row in table.findall('.//w:tr', namespaces):
                        for cell in row.findall('.//w:tc', namespaces):
                            cell_text = []
                            for text_elem in cell.findall('.//w:t', namespaces):
                                if text_elem.text:
                                    cell_text.append(text_elem.text)
                            
                            text = ''.join(cell_text).strip()
                            if text:
                                parts = re.split(r'[ \t]{3,}', text)
                                for p in parts:
                                    p = p.strip()
                                    if p:
                                        raw_segments.append(p)
    
    return raw_segments


def _is_tail_fragment(text):
    """判断是否像是上一条的尾部碎片"""
    text = text.strip()
    if not text:
        return True
    # 纯数字+单位+颜色词，不是独立机型
    if re.match(r'^\d+[GgTt]', text):
        return True
    if re.match(r'^\d+\s*50\d{0,2}', text):  # 5060, 5070
        return True
    if re.match(r'^(新品|黑色|白色|碳晶灰|鸽子灰|银灰|WIN\d+)', text):
        return True
    # 不包含任何前缀关键词
    for prefix in KNOWN_PREFIXES:
        if text.startswith(prefix):
            return False
    # 很短的文本视为碎片
    if len(text) < 10:
        return True
    return False


def _split_by_prefix(text):
    """按已知机型前缀分割一段文本为多条记录"""
    # 找到所有前缀出现位置
    positions = []
    for prefix in KNOWN_PREFIXES:
        start = 0
        while True:
            idx = text.find(prefix, start)
            if idx == -1:
                break
            # 检查是否是单词边界
            if idx > 0 and (text[idx-1].isalnum() or text[idx-1] in '））'):
                start = idx + 1
                continue
            positions.append((idx, len(prefix)))
            start = idx + 1

    if not positions:
        return [text]

    # 按位置排序
    positions.sort(key=lambda x: x[0])

    # 按前缀位置分割
    parts = []
    for i, (pos, _) in enumerate(positions):
        if i == 0:
            continue
        # 从前一个前缀到当前前缀
        prev_pos = positions[i-1][0]
        part = text[prev_pos:pos].strip()
        if part:
            parts.append(part)

    # 最后一个
    last_pos = positions[-1][0]
    tail = text[last_pos:].strip()
    if tail:
        parts.append(tail)

    return parts


def _parse_single_record(text):
    """解析单条记录，提取各字段。"""
    raw = text.strip()
    if not raw:
        return None

    result = {"raw": raw}

    # ---- 1. 提取系列名 ----
    series = None
    for prefix in KNOWN_PREFIXES:
        if raw.startswith(prefix):
            series = prefix
            break

    if series:
        result["series"] = series
        remainder = raw[len(series):].strip()
    else:
        # 尝试通用匹配：第一串连续字母数字+短横
        m = re.match(r'^([\u4e00-\u9fa5A-Za-z][\u4e00-\u9fa5A-Za-z0-9\-]*)', raw)
        if m and len(m.group(1)) >= 2:
            result["series"] = m.group(1)
            remainder = raw[len(m.group(1)):].strip()
        else:
            result["series"] = "未识别"
            remainder = raw

    # ---- 2. 清理 series 中的"联想来酷"前缀 ----
    result["series"] = result["series"].replace("联想来酷", "").strip()
    if not result["series"] or result["series"] == "未识别":
        # 如果去掉联想来酷后空了，尝试从 remainder 再取
        m = re.match(r'^([\u4e00-\u9fa5A-Za-z0-9\-]+)', remainder)
        if m:
            result["series"] = m.group(1)
            remainder = remainder[len(m.group(1)):].strip()

    if not result.get("series"):
        result["series"] = "未识别"
        return None

    # ---- 预处理：将斜杠分隔的格式（/32/1T）转为空格 ----
    # 小新系列的典型格式: ultra5-225/32/1T → ultra5-225 32 1T
    remainder = re.sub(r'/(\d+)/(\d+)T', r' \1 \2T', remainder)
    remainder = re.sub(r'/(\d+)/(\d+)G', r' \1 \2G', remainder)
    remainder = re.sub(r'/(\d+)T', r' \1T', remainder)
    remainder = re.sub(r'/(\d+)G', r' \1G', remainder)

    # ---- 3. 提取 CPU ----
    cpu_m = CPU_PATTERN.search(remainder)
    if cpu_m:
        result["cpu"] = cpu_m.group(0).strip().upper()
        remainder = remainder[:cpu_m.start()] + remainder[cpu_m.end():]
    else:
        # 可能 CPU 被包含在 series 名里（如 ultra5-225）
        cpu_in_s = CPU_PATTERN.search(result["series"])
        if cpu_in_s:
            result["cpu"] = cpu_in_s.group(0).strip().upper()
            result["series"] = result["series"].replace(cpu_in_s.group(0), "").strip().rstrip("- ")
            # 重新从完整文本获取 remainder
            series_clean = result["series"]
            if raw.startswith(series_clean):
                remainder = raw[len(series_clean):].strip()
            else:
                remainder = raw.replace(series_clean, "", 1).strip()
            # 重新对 remainder 做斜杠预处理
            remainder = re.sub(r'/(\d+)/(\d+)T', r' \1 \2T', remainder)
            remainder = re.sub(r'/(\d+)/(\d+)G', r' \1 \2G', remainder)
            remainder = re.sub(r'/(\d+)T', r' \1T', remainder)
            remainder = re.sub(r'/(\d+)G', r' \1G', remainder)
            # 再搜一次 CPU
            cpu_m = CPU_PATTERN.search(remainder)
            if cpu_m:
                result["cpu"] = cpu_m.group(0).strip().upper()
                remainder = remainder[:cpu_m.start()] + remainder[cpu_m.end():]

    # ---- 4. 提取内存 (XG 或 XGB) ----
    # 跳过显存 8G（紧跟 5060/5070 的通常是显存）
    # 先标记并移除显存
    remainder = re.sub(r'\b8G\s+(5060|5070|4060|4070)', r'VRAM_\1', remainder)
    remainder = re.sub(r'\b(8G|16G)\s+(5060|5070|4060|4070)', lambda m: f'VRAM_{m.group(2)}' if m.group(1) == '8G' else m.group(0), remainder)

    ram_m = re.search(r'\b(\d+)\s*[Gg][Bb]?\b', remainder)
    if ram_m:
        result["ram"] = ram_m.group(1) + "G"
        remainder = remainder.replace(ram_m.group(0), "", 1)
    else:
        ram_m = re.search(r'\b(\d+)\s*G\b', remainder)
        if ram_m:
            result["ram"] = ram_m.group(1) + "G"
            remainder = remainder.replace(ram_m.group(0), "", 1)

    # ---- 5. 提取硬盘 (TG/TB 或 G/GB) ----
    storage_m = re.search(r'\b(\d+)\s*[Tt][Bb]?\b', remainder)
    if storage_m:
        result["storage"] = storage_m.group(1) + "T"
        remainder = remainder.replace(storage_m.group(0), "", 1)
    else:
        storage_m = re.search(r'\b(\d+)\s*[Tt]\b', remainder)
        if storage_m:
            result["storage"] = storage_m.group(1) + "T"
            remainder = remainder.replace(storage_m.group(0), "", 1)
        else:
            storage_m = re.search(r'\b(\d+)\s*[Gg][Bb]?\b', remainder)
            if storage_m:
                result["storage"] = storage_m.group(1) + "G"
                remainder = remainder.replace(storage_m.group(0), "", 1)

    # ---- 6. 提取显卡 ----
    gpu_m = re.search(r'(\d+)\s*([45]0\d{1,2})', remainder)
    if gpu_m:
        result["gpu"] = gpu_m.group(1) + gpu_m.group(2)
        remainder = remainder.replace(gpu_m.group(0), "", 1)
    elif re.search(r'(集成|集显)', remainder):
        result["gpu"] = "集成"
        remainder = re.sub(r'(集成|集显)', "", remainder)
    else:
        gpu_m = re.search(r'(RTX\s*\d+|GTX\s*\d+)', remainder, re.IGNORECASE)
        if gpu_m:
            result["gpu"] = gpu_m.group(0).strip()
            remainder = remainder.replace(gpu_m.group(0), "", 1)

    # ---- 7. 提取屏幕尺寸 ----
    screen_m = re.search(r'\b(\d{2}(?:\.\d)?)\s*(?:英寸|寸|屏)?(?:\s|$)', remainder)
    if screen_m:
        val = screen_m.group(1)
        if 10 <= float(val) <= 18:  # 合理的笔记本屏幕尺寸
            result["screen"] = val
            remainder = remainder.replace(screen_m.group(0), "", 1)

    # ---- 8. 剩余作为备注 ----
    remainder = remainder.strip()
    remainder = re.sub(r'\s+', ' ', remainder)
    note = remainder.strip()
    # 过滤掉纯数字/纯标点的残余
    if note and not re.match(r'^[\s\-_\.\,\/]+$', note) and len(note) < 40:
        result["note"] = note
    elif note and len(note) >= 40:
        result["note"] = note[:40]

    return result


def preview_parse(filepath):
    """预览解析结果，返回表格数据用于展示"""
    products = parse_word_pricelist(filepath)
    rows = []
    for p in products:
        rows.append([
            p.get("series", ""),
            p.get("cpu", ""),
            p.get("ram", ""),
            p.get("storage", ""),
            p.get("gpu", ""),
            p.get("screen", ""),
            p.get("note", ""),
        ])
    return rows, products