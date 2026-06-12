"""
联想电脑配置提取脚本
从Word文档中提取联想电脑配置，生成规整的Word文档
支持WPS和微软Office格式
"""

import sys
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

# 联想机型前缀列表
LENOVO_PREFIXES = [
    "小新Pro16GT-ultra", "小新Pro16GT-", "小新Pro16C-", "小新Pro16-",
    "小新Pro14C-", "小新Pro14GT-", "小新Pro14-",
    "小新SE-16C-", "小新SE-14C-", "小新SE-16", "小新SE-14",
    "小新SE-16C", "小新SE-14C",
    "小新Air14-", "小新Air15-", "小新Air14", "小新Air15",
    "Y9000P-16", "Y9000P", "Y7000P", "Y7000",
    "R9000P", "R9000", "r9000p", "r9000",
    "拯救者", "LEGION",
    "ThinkPad", "ThinkBook", "Think",
    "来酷", "ideapad", "IdeaPad",
]

# CPU正则模式
CPU_PATTERN = re.compile(
    r'(I[3579][-\s]?\d{3,5}[A-Za-z]*|R[579][-\s]?\d{3,4}[A-Za-z]*'
    r'|U\d+[-\s]?\d*[A-Za-z]*|ultra\d+[-\s]?\d*[A-Za-z]*)',
    re.IGNORECASE
)


def extract_text_from_docx(filepath):
    """从Word文档提取文本，支持WPS和微软Office"""
    raw_segments = []
    
    # 方法1: 使用python-docx
    try:
        doc = Document(filepath)
        
        # 从段落提取
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts = re.split(r'[ \t]{3,}', text)
                for p in parts:
                    p = p.strip()
                    if p and len(p) > 3:
                        raw_segments.append(p)
        
        # 从表格提取
        if not raw_segments and doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            row_text.append(text)
                    if row_text:
                        combined = '    '.join(row_text)
                        raw_segments.append(combined)
        
        if raw_segments:
            return raw_segments
    except Exception as e:
        print(f"python-docx解析失败: {e}")
    
    # 方法2: 直接解析XML（兼容WPS）
    try:
        with zipfile.ZipFile(filepath, 'r') as zf:
            if 'word/document.xml' not in zf.namelist():
                raise Exception("无法找到document.xml")
            
            with zf.open('word/document.xml') as doc_xml:
                content = doc_xml.read()
                root = ET.fromstring(content)
                
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'wps': 'http://schemas.wps.cn/officeDocument/2006/main'
                }
                
                # 从段落提取
                for para in root.findall('.//w:p', namespaces):
                    para_text = []
                    for text_elem in para.findall('.//w:t', namespaces):
                        if text_elem.text:
                            para_text.append(text_elem.text)
                    
                    text = ''.join(para_text).strip()
                    if text and len(text) > 3:
                        parts = re.split(r'[ \t]{3,}', text)
                        for p in parts:
                            p = p.strip()
                            if p:
                                raw_segments.append(p)
                
                # 从表格提取
                if not raw_segments:
                    for table in root.findall('.//w:tbl', namespaces):
                        for row in table.findall('.//w:tr', namespaces):
                            row_text = []
                            for cell in row.findall('.//w:tc', namespaces):
                                cell_text = []
                                for text_elem in cell.findall('.//w:t', namespaces):
                                    if text_elem.text:
                                        cell_text.append(text_elem.text)
                                if cell_text:
                                    row_text.append(''.join(cell_text))
                            if row_text:
                                combined = '    '.join(row_text)
                                raw_segments.append(combined)
    except Exception as e:
        print(f"XML解析也失败: {e}")
    
    return raw_segments


def is_lenovo_product(text):
    """检查文本是否包含联想产品"""
    for prefix in LENOVO_PREFIXES:
        if prefix.lower() in text.lower():
            return True
    return False


def parse_product_info(text):
    """解析产品信息，提取各个参数"""
    result = {
        'series': '',
        'cpu': '',
        'ram': '',
        'storage': '',
        'gpu': '',
        'screen': '',
        'note': ''
    }
    
    raw = text.strip()
    if not raw:
        return None
    
    # 1. 提取系列名
    for prefix in LENOVO_PREFIXES:
        if prefix.lower() in raw.lower():
            result['series'] = prefix
            break
    
    if not result['series']:
        # 尝试通用匹配
        m = re.match(r'^([\u4e00-\u9fa5A-Za-z][\u4e00-\u9fa5A-Za-z0-9\-]*)', raw)
        if m and len(m.group(1)) >= 2:
            result['series'] = m.group(1)
    
    if not result['series']:
        return None
    
    # 清理系列名
    result['series'] = result['series'].replace("联想", "").strip()
    
    # 2. 提取CPU
    cpu_m = CPU_PATTERN.search(raw)
    if cpu_m:
        result['cpu'] = cpu_m.group(0).strip().upper()
    
    # 3. 提取内存
    remainder = re.sub(r'\b8G\s+(5060|5070|4060|4070)', 'VRAM', raw)
    ram_m = re.search(r'\b(\d+)\s*[Gg][Bb]?\b', remainder)
    if ram_m:
        result['ram'] = ram_m.group(1) + "G"
    
    # 4. 提取硬盘
    storage_m = re.search(r'\b(\d+)\s*[Tt][Bb]?\b', remainder)
    if storage_m:
        result['storage'] = storage_m.group(1) + "T"
    else:
        storage_m = re.search(r'\b(\d+)\s*[Gg][Bb]?\b', remainder)
        if storage_m:
            result['storage'] = storage_m.group(1) + "G"
    
    # 5. 提取显卡
    gpu_m = re.search(r'(\d+)\s*([45]0\d{1,2})', remainder)
    if gpu_m:
        result['gpu'] = gpu_m.group(1) + gpu_m.group(2)
    elif re.search(r'(RTX\s*\d+|GTX\s*\d+)', remainder, re.IGNORECASE):
        gpu_m = re.search(r'(RTX\s*\d+|GTX\s*\d+)', remainder, re.IGNORECASE)
        result['gpu'] = gpu_m.group(0).strip()
    elif re.search(r'(集成|集显)', remainder):
        result['gpu'] = "集成"
    
    # 6. 提取屏幕尺寸
    screen_m = re.search(r'\b(\d{2}(?:\.\d)?)\s*(?:英寸|寸|屏)?', remainder)
    if screen_m:
        val = screen_m.group(1)
        if 10 <= float(val) <= 18:
            result['screen'] = val
    
    # 7. 提取备注
    note_parts = []
    
    # 提取颜色信息
    colors = ['碳晶灰', '鸽子灰', '银灰', '黑色', '白色', '蓝色', '银色', '深空灰']
    for color in colors:
        if color in raw:
            note_parts.append(color)
            break
    
    # 提取其他信息
    if '新品' in raw:
        note_parts.append('新品')
    if 'Win11' in raw or 'win11' in raw:
        note_parts.append('Win11')
    
    if note_parts:
        result['note'] = ' '.join(note_parts)
    
    return result


def create_output_document(products, output_path):
    """创建规整的Word文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('联想电脑配置清单', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加生成信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_run = info_para.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    info_run.font.size = Pt(9)
    info_run.font.italic = True
    
    doc.add_paragraph()
    
    # 添加说明
    intro = doc.add_paragraph()
    intro_run = intro.add_run('说明: 本文档按照调货助手数据库格式生成，可直接导入使用。')
    intro_run.font.size = Pt(10)
    intro_run.font.italic = True
    
    doc.add_paragraph()
    
    # 创建表格
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头
    header_cells = table.rows[0].cells
    headers = ['系列', 'CPU', '内存', '硬盘', '显卡', '屏幕', '备注']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # 设置表头样式
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置列宽
    for i, width in enumerate([Cm(4), Cm(3.5), Cm(2), Cm(2), Cm(2), Cm(1.5), Cm(3)]):
        for cell in table.columns[i].cells:
            cell.width = width
    
    # 填充数据
    for product in products:
        row_cells = table.add_row().cells
        row_cells[0].text = product.get('series', '')
        row_cells[1].text = product.get('cpu', '')
        row_cells[2].text = product.get('ram', '')
        row_cells[3].text = product.get('storage', '')
        row_cells[4].text = product.get('gpu', '')
        row_cells[5].text = product.get('screen', '')
        row_cells[6].text = product.get('note', '')
        
        # 格式化单元格
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # 添加统计信息
    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.LEFT
    stats_run = stats.add_run(f'共提取 {len(products)} 条联想电脑配置')
    stats_run.font.size = Pt(11)
    stats_run.font.bold = True
    
    # 添加第二个表格：导入模板格式
    doc.add_page_break()
    doc.add_heading('调货助手导入模板', 1)
    
    template_intro = doc.add_paragraph()
    template_intro_run = template_intro.add_run(
        '以下是专门为调货助手设计的导入模板格式。如果您使用的是旧版本软件或导入功能有问题，'
        '可以使用下方的简化格式。'
    )
    template_intro_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # 创建简化表格（仅包含核心字段）
    simple_table = doc.add_table(rows=1, cols=4)
    simple_table.style = 'Table Grid'
    
    # 表头
    simple_headers = ['系列', 'CPU', '内存/硬盘', '显卡/屏幕']
    for i, header in enumerate(simple_table.rows[0].cells):
        simple_table.rows[0].cells[i].text = header
        for paragraph in simple_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 添加示例数据
    sample_data = [
        ['小新Pro16', 'Ultra5-125H', '32GB/1TB', 'RTX4060/16英寸'],
        ['Y9000P', 'i7-14700HX', '32GB/1TB', 'RTX4070/16英寸'],
        ['ThinkBook 14+', 'Ultra5-125H', '16GB/512GB', '集成/14英寸'],
    ]
    
    for row_data in sample_data:
        row = simple_table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
    
    # 保存文档
    doc.save(output_path)
    return output_path


def extract_from_word(input_path, output_path):
    """主函数：从Word文档提取联想电脑配置"""
    print("=" * 60)
    print("联想电脑配置提取工具")
    print("=" * 60)
    print(f"\n输入文件: {input_path}")
    
    if not os.path.exists(input_path):
        print(f"错误: 文件不存在 - {input_path}")
        return None
    
    # 1. 提取文本
    print("\n[1/4] 提取文本内容...")
    raw_segments = extract_text_from_docx(input_path)
    print(f"  - 提取到 {len(raw_segments)} 个文本段")
    
    if not raw_segments:
        print("  警告: 未能提取到任何文本")
        return None
    
    # 2. 过滤联想产品
    print("\n[2/4] 过滤联想产品...")
    lenovo_segments = [seg for seg in raw_segments if is_lenovo_product(seg)]
    print(f"  - 找到 {len(lenovo_segments)} 条联想相关记录")
    
    # 3. 解析产品信息
    print("\n[3/4] 解析产品参数...")
    products = []
    for seg in lenovo_segments:
        product = parse_product_info(seg)
        if product and product['series'] and product['cpu']:
            products.append(product)
    
    print(f"  - 成功解析 {len(products)} 条产品信息")
    
    # 去重
    unique_products = []
    seen = set()
    for p in products:
        key = (p['series'], p['cpu'], p['ram'], p['storage'])
        if key not in seen:
            seen.add(key)
            unique_products.append(p)
    
    print(f"  - 去重后 {len(unique_products)} 条唯一记录")
    
    if not unique_products:
        print("\n错误: 未能解析出有效的联想电脑配置")
        return None
    
    # 4. 生成Word文档
    print("\n[4/4] 生成规整Word文档...")
    output_file = create_output_document(unique_products, output_path)
    print(f"  - 已生成: {output_file}")
    
    return output_file


def main():
    """命令行入口"""
    if len(sys.argv) < 3:
        print("\n使用方法:")
        print("  python extract_lenovo_config.py <输入Word文件> <输出Word文件>")
        print("\n示例:")
        print('  python extract_lenovo_config.py "价格表.docx" "联想配置清单.docx"')
        print("\n支持的格式:")
        print("  - 微软Office Word文档 (.docx)")
        print("  - WPS Word文档 (.docx)")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    result = extract_from_word(input_file, output_file)
    
    if result:
        print("\n" + "=" * 60)
        print("处理完成!")
        print("=" * 60)
        print(f"\n输出文件: {result}")
        print(f"提取记录: {len(open(result, 'rb').read())} 条")
    else:
        print("\n处理失败，请检查输入文件格式")
        sys.exit(1)


if __name__ == "__main__":
    main()
