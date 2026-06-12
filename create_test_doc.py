"""
创建测试用的Word文档
"""
from docx import Document
from docx.shared import Pt
import os

# 创建测试文档
doc = Document()

# 添加标题
doc.add_heading('测试价格表', 0)

# 添加内容
products = [
    "小新Pro16 i7-13700H 32G 1T RTX4060 16英寸",
    "Y9000P i7-14700HX 32G 1T RTX4070 16英寸",
    "ThinkBook 14+ Ultra5-125H 16G 512G 集成 14英寸",
    "小新Pro14 Ultra7-155H 32G 1T RTX4050 14英寸 碳晶灰",
    "R9000P R9-7945HX 32G 1T RTX4060 16英寸 新品",
]

doc.add_paragraph('机型列表：')
for i, product in enumerate(products, 1):
    doc.add_paragraph(f"{i}. {product}")

# 保存
test_file = "d:\\OH-workspace\\diaohuo-assistant\\测试价格表.docx"
doc.save(test_file)
print(f"测试文档已创建: {test_file}")
