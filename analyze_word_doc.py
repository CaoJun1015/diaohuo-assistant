"""
分析Word文档格式差异
支持WPS和微软Office格式
"""
import sys
import os

def analyze_doc_file(filepath):
    """分析Word文档的内部结构"""
    print("=" * 60)
    print("Word文档结构分析")
    print("=" * 60)
    print(f"\n文件路径: {filepath}")
    
    if not os.path.exists(filepath):
        print("错误: 文件不存在!")
        return
    
    # 检查文件大小
    file_size = os.path.getsize(filepath)
    print(f"文件大小: {file_size/1024:.2f} KB")
    
    # 检查文件扩展名
    ext = os.path.splitext(filepath)[1].lower()
    print(f"文件扩展名: {ext}")
    
    # 尝试不同的解析方法
    print("\n" + "=" * 60)
    print("尝试解析文档...")
    print("=" * 60)
    
    # 方法1: 使用python-docx
    print("\n[方法1] 使用python-docx库:")
    try:
        from docx import Document
        doc = Document(filepath)
        
        print(f"段落数量: {len(doc.paragraphs)}")
        print(f"表格数量: {len(doc.tables)}")
        
        # 显示前几个段落
        print("\n前5个段落内容:")
        for i, para in enumerate(doc.paragraphs[:5]):
            text = para.text.strip()
            if text:
                print(f"  [{i+1}] {text[:80]}...")
        
        # 检查表格
        if doc.tables:
            print(f"\n第一个表格信息:")
            table = doc.tables[0]
            print(f"  行数: {len(table.rows)}")
            print(f"  列数: {len(table.columns)}")
            if len(table.rows) > 0:
                print(f"  第一行内容: {[cell.text[:20] for cell in table.rows[0].cells]}")
        
        print("\npython-docx解析成功!")
        
    except Exception as e:
        print(f"python-docx解析失败: {str(e)}")
        print(f"错误类型: {type(e).__name__}")
    
    # 方法2: 检查ZIP内部结构
    print("\n[方法2] 检查ZIP内部结构:")
    try:
        import zipfile
        with zipfile.ZipFile(filepath, 'r') as zf:
            print(f"ZIP文件列表:")
            for name in zf.namelist()[:10]:
                print(f"  - {name}")
            
            # 检查是否有WPS特有文件
            wps_files = [f for f in zf.namelist() if 'wps' in f.lower()]
            if wps_files:
                print(f"\n发现WPS特有文件:")
                for f in wps_files:
                    print(f"  - {f}")
            
            # 检查document.xml
            if 'word/document.xml' in zf.namelist():
                print("\n检查document.xml:")
                with zf.open('word/document.xml') as doc_xml:
                    content = doc_xml.read().decode('utf-8', errors='ignore')
                    print(f"  文件大小: {len(content)} 字符")
                    
                    # 检查WPS特有标记
                    if 'wps' in content.lower():
                        print("  发现WPS特有标记!")
                    
                    # 显示部分内容
                    print(f"  前200字符: {content[:200]}...")
        
        print("\nZIP结构检查成功!")
        
    except Exception as e:
        print(f"ZIP结构检查失败: {str(e)}")
    
    # 方法3: 使用olefile检查(针对.doc格式)
    if ext == '.doc':
        print("\n[方法3] 检查OLE结构:")
        try:
            import olefile
            ole = olefile.OleFileIO(filepath)
            print(f"OLE流列表:")
            for stream in ole.listdir():
                print(f"  - {'/'.join(stream)}")
            ole.close()
            print("\nOLE结构检查成功!")
        except ImportError:
            print("olefile库未安装，跳过OLE检查")
        except Exception as e:
            print(f"OLE结构检查失败: {str(e)}")

def test_wps_compatibility():
    """测试WPS兼容性"""
    print("\n" + "=" * 60)
    print("WPS vs 微软Office 兼容性测试")
    print("=" * 60)
    
    print("\n主要区别:")
    print("1. 文件格式:")
    print("   - 微软Office: 严格遵循OOXML标准")
    print("   - WPS: 基本遵循OOXML，但有扩展和差异")
    
    print("\n2. 内部结构:")
    print("   - 微软Office: 标准的XML结构")
    print("   - WPS: 可能有自定义命名空间和扩展标签")
    
    print("\n3. 编码方式:")
    print("   - 微软Office: UTF-8")
    print("   - WPS: UTF-8，但可能有BOM标记")
    
    print("\n4. 表格处理:")
    print("   - 微软Office: 标准表格结构")
    print("   - WPS: 表格可能有额外的样式信息")
    
    print("\n建议解决方案:")
    print("1. 在WPS中另存为.docx格式(选择'微软Office兼容模式')")
    print("2. 使用WPS的'导出为Microsoft Word文档'功能")
    print("3. 在微软Office中打开后重新保存")
    print("4. 更新python-docx到最新版本")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        filepath = sys.argv[1]
        analyze_doc_file(filepath)
    else:
        test_wps_compatibility()
        print("\n使用方法:")
        print("  python analyze_word_doc.py <Word文档路径>")
