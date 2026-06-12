"""
联想配置提取工具 - 带错误日志版本
"""
import sys
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from PyQt6.QtWidgets import (QApplication, QMainWindow, QLabel, QVBoxLayout, 
                             QHBoxLayout, QWidget, QMessageBox, QProgressBar,
                             QFrame)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize
from PyQt6.QtGui import QFont, QDragEnterEvent, QDropEvent, QIcon
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

LENOVO_PREFIXES = [
    "小新Pro16GT-ultra", "小新Pro16GT-", "小新Pro16C-", "小新Pro16-",
    "小新Pro14C-", "小新Pro14GT-", "小新Pro14-",
    "小新SE-16C-", "小新SE-14C-", "小新SE-16", "小新SE-14",
    "小新SE-16C", "小新SE-14C",
    "小新Air14-", "小新Air15-", "小新Air14", "小新Air15",
    "Y9000P-16", "Y9000P", "Y7000P", "Y7000",
    "R9000P", "R9000", "r9000p", "r9000",
    "拯救者", "LEGION",
    "ThinkPad", "ThinkBook", "Think",
    "来酷", "ideapad", "IdeaPad",
]

CPU_PATTERN = re.compile(
    r'(I[3579][-\s]?\d{3,5}[A-Za-z]*|R[579][-\s]?\d{3,4}[A-Za-z]*'
    r'|U\d+[-\s]?\d*[A-Za-z]*|ultra\d+[-\s]?\d*[A-Za-z]*)',
    re.IGNORECASE
)

LOG_FILE = "extract_log.txt"

def log(message):
    """写入日志文件"""
    try:
        with open(LOG_FILE, 'a', encoding='utf-8') as f:
            f.write(f"[{datetime.now().strftime('%H:%M:%S')}] {message}\n")
    except:
        pass

class ExtractThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(bool, str, str)
    
    def __init__(self, input_file):
        super().__init__()
        self.input_file = input_file
        self.products = []
        log(f"线程初始化，输入文件: {input_file}")
    
    def run(self):
        try:
            log("开始提取...")
            self.progress.emit(10)
            self.products = self.extract_products(self.input_file)
            self.progress.emit(50)
            log(f"提取完成，产品数量: {len(self.products)}")
            
            if not self.products:
                log("未找到产品")
                self.finished.emit(False, "未找到联想电脑配置", "")
                return
            
            log("开始生成输出文件...")
            output_file = self.create_output(self.products)
            self.progress.emit(100)
            log(f"输出文件生成成功: {output_file}")
            self.finished.emit(True, f"成功提取 {len(self.products)} 条配置", output_file)
            
        except Exception as e:
            import traceback
            log(f"错误: {type(e).__name__}: {e}")
            log(traceback.format_exc())
            self.finished.emit(False, f"处理失败: {str(e)}", "")
    
    def extract_products(self, filepath):
        log(f"extract_products: {filepath}")
        raw_segments = []
        
        # 方法1: python-docx
        try:
            log("尝试使用python-docx解析...")
            doc = Document(filepath)
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts = re.split(r'[ \t]{3,}', text)
                    for p in parts:
                        p = p.strip()
                        if p and len(p) > 3:
                            raw_segments.append(p)
            
            if not raw_segments and doc.tables:
                log("从表格提取...")
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text:
                                row_text.append(text)
                        if row_text:
                            combined = '    '.join(row_text)
                            raw_segments.append(combined)
            
            log(f"python-docx提取到 {len(raw_segments)} 个段落")
            
        except Exception as e:
            log(f"python-docx失败: {e}")
        
        # 方法2: XML直接解析
        if not raw_segments:
            try:
                log("尝试XML解析...")
                with zipfile.ZipFile(filepath, 'r') as zf:
                    if 'word/document.xml' not in zf.namelist():
                        log("未找到document.xml")
                        return []
                    
                    with zf.open('word/document.xml') as doc_xml:
                        content = doc_xml.read()
                        root = ET.fromstring(content)
                        
                        namespaces = {
                            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                        }
                        
                        for para in root.findall('.//w:p', namespaces):
                            para_text = []
                            for text_elem in para.findall('.//w:t', namespaces):
                                if text_elem.text:
                                    para_text.append(text_elem.text)
                            
                            text = ''.join(para_text).strip()
                            if text and len(text) > 3:
                                parts = re.split(r'[ \t]{3,}', text)
                                for p in parts:
                                    p = p.strip()
                                    if p:
                                        raw_segments.append(p)
                        
                        log(f"XML提取到 {len(raw_segments)} 个段落")
                        
            except Exception as e:
                log(f"XML解析失败: {e}")
        
        return self.parse_products(raw_segments)
    
    def parse_products(self, raw_segments):
        log(f"parse_products: {len(raw_segments)} 个段落")
        
        lenovo_segments = [seg for seg in raw_segments if any(prefix.lower() in seg.lower() for prefix in LENOVO_PREFIXES)]
        log(f"找到 {len(lenovo_segments)} 条联想相关记录")
        
        products = []
        for seg in lenovo_segments:
            product = self.parse_single(seg)
            if product and product['series'] and product['cpu']:
                products.append(product)
        
        # 去重
        unique_products = []
        seen = set()
        for p in products:
            key = (p['series'], p['cpu'], p['ram'], p['storage'])
            if key not in seen:
                seen.add(key)
                unique_products.append(p)
        
        log(f"去重后 {len(unique_products)} 条产品")
        return unique_products
    
    def parse_single(self, text):
        result = {'series': '', 'cpu': '', 'ram': '', 'storage': '', 'gpu': '', 'screen': '', 'note': ''}
        
        raw = text.strip()
        if not raw:
            return None
        
        for prefix in LENOVO_PREFIXES:
            if prefix.lower() in raw.lower():
                result['series'] = prefix
                break
        
        if not result['series']:
            m = re.match(r'^([\u4e00-\u9fa5A-Za-z][\u4e00-\u9fa5A-Za-z0-9\-]*)', raw)
            if m and len(m.group(1)) >= 2:
                result['series'] = m.group(1)
        
        if not result['series']:
            return None
        
        result['series'] = result['series'].replace("联想", "").strip()
        
        cpu_m = CPU_PATTERN.search(raw)
        if cpu_m:
            result['cpu'] = cpu_m.group(0).strip().upper()
        
        remainder = re.sub(r'\b8G\s+(5060|5070|4060|4070)', 'VRAM', raw)
        ram_m = re.search(r'\b(\d+)\s*[Gg][Bb]?\b', remainder)
        if ram_m:
            result['ram'] = ram_m.group(1) + "G"
        
        storage_m = re.search(r'\b(\d+)\s*[Tt][Bb]?\b', remainder)
        if storage_m:
            result['storage'] = storage_m.group(1) + "T"
        else:
            storage_m = re.search(r'\b(\d+)\s*[Gg][Bb]?\b', remainder)
            if storage_m:
                result['storage'] = storage_m.group(1) + "G"
        
        gpu_m = re.search(r'(\d+)\s*([45]0\d{1,2})', remainder)
        if gpu_m:
            result['gpu'] = gpu_m.group(1) + gpu_m.group(2)
        elif re.search(r'(RTX\s*\d+|GTX\s*\d+)', remainder, re.IGNORECASE):
            gpu_m = re.search(r'(RTX\s*\d+|GTX\s*\d+)', remainder, re.IGNORECASE)
            result['gpu'] = gpu_m.group(0).strip()
        elif re.search(r'(集成|集显)', remainder):
            result['gpu'] = "集成"
        
        screen_m = re.search(r'\b(\d{2}(?:\.\d)?)\s*(?:英寸|寸|屏)?', remainder)
        if screen_m:
            val = screen_m.group(1)
            if 10 <= float(val) <= 18:
                result['screen'] = val
        
        note_parts = []
        colors = ['碳晶灰', '鸽子灰', '银灰', '黑色', '白色', '蓝色', '银色', '深空灰']
        for color in colors:
            if color in raw:
                note_parts.append(color)
                break
        
        if '新品' in raw:
            note_parts.append('新品')
        if 'Win11' in raw or 'win11' in raw:
            note_parts.append('Win11')
        
        if note_parts:
            result['note'] = ' '.join(note_parts)
        
        return result
    
    def create_output(self, products):
        log("create_output开始...")
        
        app_path = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(__file__)
        input_name = os.path.splitext(os.path.basename(self.input_file))[0]
        output_file = os.path.join(app_path, f"{input_name}_联想配置.docx")
        
        log(f"输出文件路径: {output_file}")
        
        try:
            doc = Document()
            log("Document创建成功")
            
            title = doc.add_heading('联想电脑配置清单', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            log("标题添加成功")
            
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            info_run = info_para.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            info_run.font.size = Pt(9)
            info_run.font.italic = True
            log("生成时间添加成功")
            
            doc.add_paragraph()
            
            intro = doc.add_paragraph()
            intro_run = intro.add_run('说明: 本文档按照调货助手数据库格式生成，可直接导入使用。')
            intro_run.font.size = Pt(10)
            intro_run.font.italic = True
            
            doc.add_paragraph()
            log("说明添加成功")
            
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            log("表格创建成功")
            
            header_cells = table.rows[0].cells
            headers = ['系列', 'CPU', '内存', '硬盘', '显卡', '屏幕', '备注']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            log("表头添加成功")
            
            for i, width in enumerate([Cm(4), Cm(3.5), Cm(2), Cm(2), Cm(2), Cm(1.5), Cm(3)]):
                for cell in table.columns[i].cells:
                    cell.width = width
            
            log(f"开始添加 {len(products)} 条产品...")
            for idx, product in enumerate(products):
                log(f"添加产品 {idx+1}/{len(products)}")
                row_cells = table.add_row().cells
                row_cells[0].text = product.get('series', '')
                row_cells[1].text = product.get('cpu', '')
                row_cells[2].text = product.get('ram', '')
                row_cells[3].text = product.get('storage', '')
                row_cells[4].text = product.get('gpu', '')
                row_cells[5].text = product.get('screen', '')
                row_cells[6].text = product.get('note', '')
                
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            doc.add_paragraph()
            
            stats = doc.add_paragraph()
            stats.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats_run = stats.add_run(f'共提取 {len(products)} 条联想电脑配置')
            stats_run.font.size = Pt(11)
            stats_run.font.bold = True
            
            log("保存文档...")
            doc.save(output_file)
            log("文档保存成功")
            
            return output_file
            
        except Exception as e:
            log(f"create_output失败: {type(e).__name__}: {e}")
            import traceback
            log(traceback.format_exc())
            raise


class DropArea(QFrame):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setFrameStyle(QFrame.Shape.Box | QFrame.Shadow.Raised)
        self.setLineWidth(2)
        self.setMinimumSize(400, 300)
        
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.icon_label = QLabel("📄")
        self.icon_label.setFont(QFont("Arial", 60))
        self.icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.text_label = QLabel("拖拽Word文档到此处")
        self.text_label.setFont(QFont("Microsoft YaHei", 14))
        self.text_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.hint_label = QLabel("(支持WPS和微软Office)")
        self.hint_label.setFont(QFont("Microsoft YaHei", 10))
        self.hint_label.setStyleSheet("color: gray;")
        self.hint_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.format_label = QLabel("支持格式: .docx")
        self.format_label.setFont(QFont("Microsoft YaHei", 9))
        self.format_label.setStyleSheet("color: #888;")
        self.format_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        layout.addWidget(self.icon_label)
        layout.addWidget(self.text_label)
        layout.addWidget(self.hint_label)
        layout.addWidget(self.format_label)
        
        log("DropArea初始化完成")
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()
            if urls and urls[0].toLocalFile().lower().endswith('.docx'):
                event.acceptProposedAction()
                self.setStyleSheet("border: 3px dashed #4CAF50; background-color: #E8F5E9;")
                return
        event.ignore()
    
    def dragLeaveEvent(self, event):
        self.setStyleSheet("")
    
    def dropEvent(self, event: QDropEvent):
        log("检测到文件拖拽")
        self.setStyleSheet("")
        urls = event.mimeData().urls()
        if urls:
            # 使用toLocalFile()并处理可能的编码问题
            file_path = urls[0].toLocalFile()
            log(f"原始文件路径: {repr(file_path)}")
            
            # 确保路径是正确的字符串格式
            if isinstance(file_path, bytes):
                file_path = file_path.decode('utf-8', errors='ignore')
            else:
                file_path = str(file_path)
            
            log(f"处理后文件路径: {file_path}")
            
            # 检查文件是否存在
            if not os.path.exists(file_path):
                log(f"文件不存在: {file_path}")
                QMessageBox.warning(self.parent(), "错误", "文件不存在或路径无效")
                return
            
            # 检查文件大小
            try:
                file_size = os.path.getsize(file_path)
                log(f"文件大小: {file_size} bytes")
            except Exception as e:
                log(f"无法获取文件大小: {e}")
            
            if file_path.lower().endswith('.docx'):
                log("准备调用process_file...")
                log(f"parent对象: {self.parent()}")
                log(f"parent类型: {type(self.parent())}")
                
                # 直接调用MainWindow的方法
                main_window = self.parent()
                if main_window and hasattr(main_window, 'process_file'):
                    log("调用main_window.process_file...")
                    try:
                        main_window.process_file(file_path)
                        log("process_file调用完成")
                    except Exception as e:
                        log(f"process_file调用失败: {type(e).__name__}: {e}")
                        import traceback
                        log(traceback.format_exc())
                else:
                    log("错误: parent没有process_file方法")
            else:
                log(f"不是docx文件: {file_path}")
                QMessageBox.warning(self.parent(), "错误", "请选择Word文档(.docx格式)")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("联想配置提取工具")
        self.setMinimumSize(500, 400)
        self.thread = None
        
        log("MainWindow初始化")
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        title_label = QLabel("联想配置提取工具")
        title_label.setFont(QFont("Microsoft YaHei", 18, QFont.Weight.Bold))
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(title_label)
        
        main_layout.addSpacing(20)
        
        self.drop_area = DropArea(self)
        main_layout.addWidget(self.drop_area, 1)
        
        main_layout.addSpacing(20)
        
        status_frame = QFrame()
        status_frame.setFrameStyle(QFrame.Shape.Box | QFrame.Shadow.Sunken)
        status_layout = QVBoxLayout(status_frame)
        
        self.status_label = QLabel("状态: 等待中...")
        self.status_label.setFont(QFont("Microsoft YaHei", 10))
        status_layout.addWidget(self.status_label)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        status_layout.addWidget(self.progress_bar)
        
        self.file_label = QLabel("")
        self.file_label.setFont(QFont("Microsoft YaHei", 9))
        self.file_label.setStyleSheet("color: #666;")
        self.file_label.setWordWrap(True)
        self.file_label.setVisible(False)
        status_layout.addWidget(self.file_label)
        
        main_layout.addWidget(status_frame)
        
        log("MainWindow初始化完成")
    
    def process_file(self, file_path):
        log(f"开始处理文件: {file_path}")
        
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                log(f"文件不存在")
                QMessageBox.warning(self, "错误", "文件不存在")
                return
            
            log("准备创建线程...")
            self.file_label.setText(f"已接收: {os.path.basename(file_path)}")
            self.file_label.setVisible(True)
            self.status_label.setText("状态: 正在处理...")
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            
            log("创建ExtractThread实例...")
            self.thread = ExtractThread(file_path)
            
            log("连接信号...")
            self.thread.progress.connect(self.on_progress)
            self.thread.finished.connect(self.on_finished)
            
            log("启动线程...")
            self.thread.start()
            log("线程已启动")
            
        except Exception as e:
            log(f"创建线程失败: {type(e).__name__}: {e}")
            import traceback
            log(traceback.format_exc())
            self.progress_bar.setVisible(False)
            self.status_label.setText(f"失败: {str(e)}")
            self.status_label.setStyleSheet("color: #F44336;")
            QMessageBox.critical(self, "严重错误", f"无法启动处理:\n{str(e)}")
    
    def on_progress(self, value):
        self.progress_bar.setValue(value)
    
    def on_finished(self, success, message, output_file):
        log(f"on_finished: success={success}, message={message}, output={output_file}")
        self.progress_bar.setVisible(False)
        
        if success:
            self.status_label.setText(f"成功: {message}")
            self.status_label.setStyleSheet("color: #4CAF50; font-weight: bold;")
            
            msg = QMessageBox(self)
            msg.setWindowTitle("处理完成")
            msg.setText(f"成功提取联想电脑配置！")
            msg.setInformativeText(f"输出文件:\n{output_file}")
            msg.setIcon(QMessageBox.Icon.Information)
            
            open_btn = msg.addButton("打开文件", QMessageBox.ButtonRole.AcceptRole)
            msg.addButton("确定", QMessageBox.ButtonRole.RejectRole)
            
            msg.exec()
            
            if msg.clickedButton() == open_btn:
                os.startfile(output_file)
            
            self.status_label.setText("状态: 等待中...")
            self.status_label.setStyleSheet("")
            self.file_label.setVisible(False)
        else:
            self.status_label.setText(f"失败: {message}")
            self.status_label.setStyleSheet("color: #F44336;")
            
            QMessageBox.warning(self, "处理失败", message + "\n\n详细信息已保存到日志文件")


def main():
    log("程序启动")
    app = QApplication(sys.argv)
    app.setFont(QFont("Microsoft YaHei", 10))
    
    window = MainWindow()
    window.show()
    
    log("窗口显示")
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
